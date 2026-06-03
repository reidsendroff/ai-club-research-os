from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = r"C:\Projects\AI_Club_Brain\AI_Club_Leadership_Accountability_Audit_and_Rewrite.docx"


# COLORS1= {
#     "bob", RGBColor(255, 255, 255),
#     "bat", RGBColor(0,0,0),
# }

COLORS = {
    "ink": RGBColor(24, 24, 27),
    "muted": RGBColor(82, 82, 91),
    "purple": RGBColor(99, 102, 241),
    "teal": RGBColor(13, 148, 136),
    "amber": RGBColor(180, 83, 9),
    "red": RGBColor(185, 28, 28),
    "green": RGBColor(21, 128, 61),
    "blue": RGBColor(37, 99, 235),
    "light_purple": "EEF2FF",
    "light_teal": "F0FDFA",
    "light_amber": "FFFBEB",
    "light_red": "FEF2F2",
    "light_gray": "F4F4F5",
    "white": "FFFFFF",
    "border": "D4D4D8",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D4D4D8", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_run(run, bold=False, italic=False, color=None, size=None):
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)


def para(doc, text="", style=None, color=None, bold=False, italic=False, size=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    if text:
        r = p.add_run(text)
        style_run(r, bold=bold, italic=italic, color=color, size=size)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    style_run(r, bold=True, color=COLORS["purple"], size=18)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    style_run(r, bold=True, color=COLORS["ink"], size=13)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    style_run(r, color=color or COLORS["ink"], size=10)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    style_run(r, color=COLORS["ink"], size=10)
    return p


def callout(doc, label, text, fill="F0FDFA", accent=COLORS["teal"]):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table, "FFFFFF", "0")
    cells = table.rows[0].cells
    set_cell_width(cells[0], 160)
    set_cell_width(cells[1], 9200)
    set_cell_shading(cells[0], "%02X%02X%02X" % (accent[0], accent[1], accent[2]))
    set_cell_shading(cells[1], fill)
    for c in cells:
        set_cell_margins(c, 120, 120, 120, 120)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    style_run(r, bold=True, color=accent, size=10)
    r = p.add_run(text)
    style_run(r, color=COLORS["ink"], size=10)
    para(doc, "", after=3)
    return table


def simple_table(doc, headers, rows, widths=None, header_fill="EEF2FF"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 9360)
    set_table_borders(table)
    hdr = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i], 110, 120, 110, 120)
        if widths:
            set_cell_width(hdr[i], widths[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, bold=True, color=COLORS["ink"], size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if widths:
                set_cell_width(cells[i], widths[i])
            set_cell_margins(cells[i], 110, 120, 110, 120)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, color=COLORS["ink"], size=8.7)
    para(doc, "", after=4)
    return table


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("NORTHERN HIGHLANDS AI CLUB")
    style_run(r, bold=True, color=COLORS["purple"], size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Leadership Accountability System")
    style_run(r, bold=True, color=COLORS["ink"], size=24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Blunt audit of the company-style proposal, plus a corrected operating model for Ben review")
    style_run(r, color=COLORS["muted"], size=12)

    simple_table(
        doc,
        ["Item", "Value"],
        [
            ["Prepared for", "Reid Sendroff and Ben Shamosh"],
            ["Prepared by", "Codex, acting from the supplied research-engineer super prompt"],
            ["Date", "May 31, 2026"],
            ["Source base", "32 Markdown files, AI_Club_June1_Meeting_Plan.docx, AI_Club_Strategic_Plan_2026.docx"],
            ["Limitation", "AI Club Leadership Google Drive link was not provided, so live Drive context was not reviewed."],
        ],
        widths=[2200, 7160],
    )

    callout(
        doc,
        "Bottom line",
        "The instinct is right. The execution Claude produced is not ready to present. The corrected version should be framed as a leadership operating system, not a student-club HR department.",
        fill=COLORS["light_amber"],
        accent=COLORS["amber"],
    )
    doc.add_page_break()


def add_super_prompt(doc):
    h1(doc, "Super Prompt Executed")
    para(doc, "This is the prompt I used as the operating instruction for the audit and rewrite.", color=COLORS["muted"])
    prompt = (
        "Act as a top research engineer in the top 0.1% of research engineers, with the judgment of an elite operator who has built high-accountability technical teams. "
        "Audit the proposed AI Club company-style leadership accountability section against all local AI Club context, the June 1 meeting plan, the strategic plan, and every Markdown file in the workspace. "
        "Be blunt. Identify every structural, operational, cultural, legal/school-context, incentive, timing, and language problem. Do not stop at criticism. Produce a corrected leadership accountability system that preserves the user's intent: leaders need written role criteria, dated deliverables, evidence-based review, a fair improvement process, and replacement when a role is not being performed. "
        "Translate corporate concepts into a student-appropriate operating system. Avoid performative HR language. Keep the bar high without creating public shame, favoritism, or advisor risk. Integrate with the existing AI Club OS: building club identity, 8-officer structure, 75% member participation floor, 90% leadership attendance expectation, bi-weekly scrums, Room 133 open lab, parent donation drive, API fund, competition calendar, and succession protocol. "
        "Return a usable document for Reid and Ben: executive verdict, precise failure list, corrected framework, role scorecards, dated implementation calendar, leadership improvement plan template, meeting-script language, and document-integration recommendations."
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(prompt)
    style_run(r, italic=True, color=COLORS["ink"], size=10)

    h2(doc, "Files Actually Reviewed")
    bullet(doc, "32 Markdown files were enumerated and reviewed for AI Club context, including AI_Club_OS, AI_Club_Docs, AI_Club.md, CLAUDE.md, MEGA_CONTEXT.md, MASTER_PROMPT.md, and June1_Meeting_SuperPrompt.md.")
    bullet(doc, "The current AI_Club_June1_Meeting_Plan.docx was extracted and reviewed, including Claude's added Part 1B.")
    bullet(doc, "AI_Club_Strategic_Plan_2026.docx was extracted and reviewed because that was the DOCX named in the user context.")
    bullet(doc, "Official Anthropic and OpenAI pages were checked only for the time-sensitive AI-tools claims in the existing June 1 plan.")
    doc.add_page_break()


def add_audit(doc):
    h1(doc, "Blunt Audit")
    callout(
        doc,
        "Verdict",
        "Claude did a decent first-pass brainstorm. It did not produce a safe operating policy. The current Part 1B is too corporate, under-specified where fairness matters, over-specified where incentives can backfire, and bolted into a meeting plan that no longer fits the clock.",
        fill=COLORS["light_red"],
        accent=COLORS["red"],
    )

    h2(doc, "What Claude Got Right")
    bullet(doc, "It understood the real problem: leadership cannot be decorative if the club is going to survive without Reid doing the work.")
    bullet(doc, "It correctly tied accountability to role-specific outputs instead of vague effort.")
    bullet(doc, "It added a recovery process before removal, which is better than surprise replacement.")
    bullet(doc, "It recognized succession as part of accountability, not a separate April-only event.")

    h2(doc, "Everything Wrong Enough to Fix Before Ben Sees It")
    issues = [
        ("Agenda math is broken", "The June 1 plan now runs about 75 minutes before real-world drift. Adding a 10-minute Part 1B without cutting Part 2 or Part 3 makes the document operationally false."),
        ("Wrong document target", "The user context names AI_Club_Strategic_Plan_2026.docx, but Claude updated AI_Club_June1_Meeting_Plan.docx. That may be useful, but it ignores the explicit referenced file."),
        ("The phrase 'run it like a company' is dangerous if used raw", "It will sound intense to Reid and Ben, but to members, parents, and an advisor it can sound like unpaid minors are being put through workplace discipline."),
        ("'PIP' language is the wrong surface language", "Use Leadership Improvement Plan internally. Do not announce 'PIP' to a high school club unless the goal is to make everyone anxious and invite adult questions."),
        ("It violates the club's own language rules", "The section uses em dashes and corporate language even though CLAUDE.md and the announcement system explicitly ban both across club output."),
        ("It creates an HR process without HR safeguards", "There is no confidentiality model, appeal path, advisor role, documentation access rule, or definition of who can see performance records."),
        ("It puts the Secretary in the wrong job", "Having the Secretary document officer performance reviews mixes neutral records with confidential personnel-style notes. Secretary can log artifacts and attendance, not private improvement plans."),
        ("The KPIs are shallow in places", "Follower growth is a bad CMO metric. It incentivizes attention hunting. The real metric is whether every meeting is promoted clearly, on time, with specific demos and useful recaps."),
        ("It does not define evidence", "Green, Yellow, and Red only work if each status maps to observable artifacts: agenda posted, attendance logged, ledger updated, blocker closed, post published."),
        ("It skips a coaching step", "A company-style PIP is usually late-stage. For a student club, the first response should be a direct reset conversation and support plan before a formal improvement plan."),
        ("The replacement standard is too abrupt", "'Two Yellows to Red' ignores illness, school load, family issues, or role mis-scoping. The process needs judgment, not a vending machine."),
        ("It creates weird rewards", "'First pick of the best projects and competition teams' undermines balanced project-team formation and risks making leadership a status economy."),
        ("The bench idea is half-right but socially clumsy", "Keeping runner-up applicants warm is smart. Calling them a bench can make members feel like backups. Use 'ready successors' or 'alternate candidates.'"),
        ("It does not hold Presidents accountable", "If the goal is systems, Reid and Ben need standards too: response time, advisor communication, admin asks, final approvals, and not filling delegated roles."),
        ("It does not integrate with the existing Leadership Failure Protocol", "AI_Club_OS already has direct conversation, written note, then step-down or reassignment. The new model should upgrade that protocol, not create a competing one."),
        ("It ignores advisor and school-policy risk", "Removal from leadership in a school club should be advisor-visible. Any formal written improvement plan should be reviewed by Mr. Novak before it starts."),
        ("It over-relies on monthly reviews", "Monthly reviews are too slow for weekly operations. Use weekly artifact checks and monthly conversations."),
        ("It underuses the existing scrum system", "Leadership accountability should be managed the same way member work is managed: visible commitments, next sprint, blocker log, evidence."),
        ("It does not define role charters", "A one-page job description is mentioned but not actually specified. The corrected system needs a required charter template."),
        ("It is too much for the June 1 meeting", "Members do not need the full policy in the room on June 1. Ben needs it. Members need the standard, the application process, and the fact that roles will be reviewed."),
    ]
    simple_table(
        doc,
        ["Problem", "Why it matters"],
        issues,
        widths=[2600, 6760],
        header_fill=COLORS["light_red"],
    )

    h2(doc, "The Clean Diagnosis")
    bullet(doc, "Claude solved the wrong layer. It wrote a motivational accountability section. The club needs an operating system.")
    bullet(doc, "The policy must be enforceable from artifacts, not from vibes.")
    bullet(doc, "The language needs to be high-standard but school-safe.")
    bullet(doc, "The full framework belongs in the Strategic Plan / AI Club OS. The June 1 meeting gets a short version.")
    doc.add_page_break()


def add_corrected_framework(doc):
    h1(doc, "Corrected Framework")
    callout(
        doc,
        "Rename it",
        "Do not call this 'Running the Club Like a Company' in the final member-facing version. Call it the Leadership Operating System. Internally, Reid and Ben can keep the company analogy.",
        fill=COLORS["light_purple"],
        accent=COLORS["purple"],
    )

    h2(doc, "Operating Principle")
    para(
        doc,
        "Leadership in the AI Club is an operating role. Every officer owns a defined function, produces visible artifacts on a schedule, gets regular feedback, and can be reassigned if the role is not being performed. The point is not punishment. The point is to protect the club from becoming dependent on one or two people.",
    )

    h2(doc, "The Four Rules")
    numbered(doc, "Every role has a charter. Mission, weekly responsibilities, recurring deadlines, required artifacts, backup owner, and what Green performance looks like.")
    numbered(doc, "Performance is based on evidence. Agendas, posts, ledgers, attendance logs, blocker logs, and competition calendars count. Good intentions do not count by themselves.")
    numbered(doc, "Feedback is routine. Officers get weekly artifact checks, monthly role reviews, and marking-period status decisions.")
    numbered(doc, "Underperformance gets support first, then reassignment if needed. Direct reset, written improvement plan, advisor-visible review, role transition if the plan is not met.")

    h2(doc, "Status System")
    simple_table(
        doc,
        ["Status", "Meaning", "Action"],
        [
            ["Green", "Role is being performed reliably. Artifacts are on time and quality is acceptable.", "Continue. Consider more ownership, public credit, or mentoring a successor."],
            ["Yellow", "One or more expected artifacts are late, missing, weak, or repeatedly need rescue.", "Direct reset conversation. Written next-step target due before the next review point."],
            ["Red", "Core role function is not happening, or the same Yellow issue repeats after reset.", "Leadership Improvement Plan. Advisor-visible. Clear deliverables and deadline."],
        ],
        widths=[1200, 4400, 3760],
    )

    h2(doc, "Process Ladder")
    simple_table(
        doc,
        ["Step", "When it happens", "What happens"],
        [
            ["1. Direct reset", "First meaningful miss", "Co-President or VP/COO names the miss, asks what is blocking it, and sets the next concrete deliverable."],
            ["2. Written reset", "Repeat miss or missed reset deliverable", "Short written note: role standard, evidence, next deliverable, due date, support offered."],
            ["3. Leadership Improvement Plan", "Core role failure or repeated Yellow", "2 meeting cycles or 2 to 3 weeks. Specific recovery deliverables. Advisor is aware before it starts."],
            ["4. Role transition", "Plan not met", "Officer is reassigned or steps down. Successor or interim owner takes the charter within one week."],
        ],
        widths=[1900, 2500, 4960],
    )
    doc.add_page_break()


def add_role_scorecards(doc):
    h1(doc, "Role Scorecards")
    para(doc, "These scorecards are intentionally artifact-based. If the artifact exists and is good enough, the role is functioning. If not, it is not.")

    roles = [
        (
            "VP / Chief Operations Officer",
            [
                ["Agenda", "Meeting agenda distributed to officers 48 hours before every meeting.", "Shared agenda timestamp."],
                ["Room and AV", "Room 133, screen, and meeting setup confirmed by lunch on meeting day.", "Setup checklist."],
                ["Scrums", "Every project team presents at least every two weeks, with time kept and commitments captured.", "Scrum tracker."],
                ["Blockers", "Team blockers logged within 24 hours and followed up within 5 school days.", "Blocker log."],
                ["Team liaison", "Each project team has a named VP/COO contact and receives at least one check-in between scrums.", "Liaison notes."],
            ],
        ),
        (
            "CMO / Social Media Manager",
            [
                ["7AM hook", "Posted or scheduled for every meeting, under the club's tone rules.", "Post link or screenshot."],
                ["1PM reminder", "Posted or scheduled with real agenda detail and Room 133 @ 2:45.", "Post link or screenshot."],
                ["You Missed It", "Posted within 3 hours of the meeting ending, based on what actually happened.", "Post link or screenshot."],
                ["Content calendar", "Next two meetings' content needs collected from VP/COO by the prior weekend.", "Content calendar."],
                ["Quality", "No generic hype, no fake polish, no off-brand sign-off, no missing demo specifics.", "Co-President review notes."],
            ],
        ),
        (
            "Treasurer",
            [
                ["Ledger", "Club ledger updated within 48 hours of every transaction.", "Ledger timestamp."],
                ["Donation drive", "September parent donation process prepared, launched, and closed with totals reported.", "Donation tracker."],
                ["API fund", "Team API funds distributed at the start of Q2 after teams are scoped.", "Distribution record."],
                ["Reimbursements", "Approved reimbursements processed within 5 school days of receipt.", "Reimbursement log."],
                ["Monthly report", "Balance, spending, and risks shared with officers monthly.", "Monthly report."],
            ],
        ),
        (
            "Secretary",
            [
                ["Attendance", "Attendance logged within 24 hours of every meeting and open lab.", "Attendance ledger."],
                ["Warnings", "Members at or below 80% receive written notice within 3 school days.", "Warning log."],
                ["Strike tracking", "Scrum absences and no-progress exceptions tracked accurately.", "Strike log."],
                ["Competition calendar", "Semester competition calendar published in September and January.", "Calendar link."],
                ["Minutes", "Meeting decisions and action items posted within 24 hours.", "Minutes doc."],
            ],
        ),
        (
            "Co-Presidents",
            [
                ["Delegation", "Do not silently absorb delegated work. If an officer misses, use the process.", "Weekly leadership notes."],
                ["Advisor relationship", "Mr. Novak gets clear asks, risk updates, and policy changes before they surprise him.", "Advisor update log."],
                ["Approvals", "External communications and expenses over $50 reviewed within 24 hours when possible.", "Approval log."],
                ["Admin asks", "Room 133, tool unblocks, parent donation process, and GitHub/Drive infrastructure moved forward.", "Admin tracker."],
                ["Succession", "Potential successors identified by February and shadowing starts by March.", "Succession tracker."],
            ],
        ),
    ]

    for role, rows in roles:
        h2(doc, role)
        simple_table(doc, ["Area", "Green standard", "Evidence"], rows, widths=[1600, 5200, 2560], header_fill=COLORS["light_teal"])


def add_improvement_template(doc):
    doc.add_page_break()
    h1(doc, "Leadership Improvement Plan Template")
    callout(
        doc,
        "Use internally",
        "For members, say 'improvement plan' or 'role reset.' Do not brand this as a PIP in announcements.",
        fill=COLORS["light_amber"],
        accent=COLORS["amber"],
    )
    simple_table(
        doc,
        ["Field", "What to write"],
        [
            ["Officer / role", "Name and exact role."],
            ["Role standard missed", "Quote the charter standard, not a vague complaint."],
            ["Evidence", "What artifact was missing, late, incomplete, or rescued by someone else."],
            ["Impact", "What broke or almost broke because of it."],
            ["Recovery deliverables", "1 to 3 concrete outputs, each with a due date."],
            ["Support", "Who will help, when they will check in, and what blocker will be removed."],
            ["Timeline", "Normally two meeting cycles or 2 to 3 weeks. Adjust for school breaks."],
            ["Review owner", "One Co-President, with Mr. Novak aware for role-removal risk."],
            ["Outcomes", "Met: return to Green or Yellow with monitoring. Not met: role reassignment or transition."],
            ["Privacy", "Stored only in the leadership folder. Not discussed publicly. Secretary logs only final role-status changes."],
        ],
        widths=[2400, 6960],
        header_fill=COLORS["light_amber"],
    )

    h2(doc, "Example")
    bullet(doc, "Role: CMO.")
    bullet(doc, "Standard missed: three-post system for every meeting.")
    bullet(doc, "Evidence: 1PM reminder missed twice in three meetings; You Missed It posted 28 hours late.")
    bullet(doc, "Recovery deliverables: publish full three-post set for the next two meetings, drafts due to Reid and Ben by Sunday 8 PM before each meeting.")
    bullet(doc, "Support: VP/COO sends agenda details by Saturday noon; Reid reviews draft by Sunday night.")
    bullet(doc, "Outcome: if both cycles hit, CMO returns to Green. If one is missed without prior communication, role transitions to alternate candidate or interim owner.")


def add_timeline_and_language(doc):
    doc.add_page_break()
    h1(doc, "Implementation Calendar")
    simple_table(
        doc,
        ["Date / window", "Action"],
        [
            ["June 1", "Announce leadership applications and the fact that leadership roles will have written expectations next year. Do not present the full accountability policy."],
            ["June 8", "Applications close. Candidates identify target role."],
            ["June 10", "Recorded demo deadline for candidates not presenting live."],
            ["June 15", "Candidate demos. Evaluate project quality plus role fit."],
            ["June 16", "Decisions announced privately first, then to the club."],
            ["June 17", "New leadership Discord/server channel opens. Each officer receives a draft charter."],
            ["August", "Reid and Ben finalize role charters, scorecards, Drive folders, templates, and advisor review."],
            ["September Week 1", "Leadership kickoff. Officers sign charters. Member onboarding and interest forms begin."],
            ["September Week 2-3", "Project teams formed. VP/COO liaisons assigned. Parent donation drive active."],
            ["October", "First full scrum cycle. API funds distributed after teams are stable. First monthly officer review."],
            ["November", "First formal status review before Congressional App Challenge deadline pressure peaks."],
            ["December", "Mid-year review: role fit, attendance, financial reconciliation, team health."],
            ["January", "Second-semester reset and new competition calendar."],
            ["February-March", "Successor identification and shadowing conversations begin."],
            ["April-May", "Officer elections, transition documents, incoming-officer shadowing."],
            ["June", "Full handoff. Old officers become available advisors, not hidden operators."],
        ],
        widths=[2200, 7160],
    )

    h1(doc, "What to Say on June 1")
    callout(
        doc,
        "Keep it short",
        "The June 1 meeting should not become a policy seminar. The message is the standard, not the machinery.",
        fill=COLORS["light_purple"],
        accent=COLORS["purple"],
    )
    para(
        doc,
        "Suggested script:",
        bold=True,
        color=COLORS["ink"],
    )
    script = (
        "Next year, leadership is going to work differently. Every officer role will have a written job description, specific weekly responsibilities, and real review points during the year. "
        "If someone takes a leadership role and the work is not getting done, we are going to address it directly, support them, and if it still does not change, move someone else into the role. "
        "That is not meant to be harsh. It is the only way this club becomes bigger than Reid and Ben. Leadership here has to mean ownership."
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(script)
    style_run(r, italic=True, color=COLORS["ink"], size=10.5)

    h1(doc, "Where This Belongs")
    simple_table(
        doc,
        ["Document", "Recommendation"],
        [
            ["AI_Club_Strategic_Plan_2026.docx", "Add a condensed Leadership Operating System section after Organizational Structure or Participation & Accountability. This is the right admin/advisor-facing home."],
            ["AI_Club_June1_Meeting_Plan.docx", "Replace the current 10-minute Part 1B with a 3 to 5 minute leadership-standard note, or cut Part 2/Part 3 enough to make the agenda real."],
            ["AI_Club_OS/02_Leadership_and_Roles.md", "Replace the existing Leadership Failure Protocol with the upgraded ladder: direct reset, written reset, improvement plan, role transition."],
            ["AI_Club_OS/04_Member_System.md", "Keep member attendance and strike policy separate from officer performance. Do not mix member discipline and officer role review."],
            ["AI_Club_OS/10_Succession_Protocol.md", "Add ready-successor logic for mid-year officer replacement."],
        ],
        widths=[3000, 6360],
        header_fill=COLORS["light_teal"],
    )


def add_ben_review(doc):
    doc.add_page_break()
    h1(doc, "Ben Review Checklist")
    para(doc, "These are the questions Reid and Ben should answer before anything is shown to members.")
    for item in [
        "Are we comfortable enforcing role transitions mid-year, or only at semester boundaries except for major misses?",
        "What role does Mr. Novak need to play before a leader is removed?",
        "Who can see improvement plans? Recommended answer: Reid, Ben, the officer, and Mr. Novak if removal is possible.",
        "Do we want 90% leadership attendance as a hard standard, or a Green standard with exceptions?",
        "Should current officers reapply through the same project-demo process, or submit role-specific evidence instead?",
        "What is the exact replacement path if no alternate candidate is ready?",
        "How do we avoid making leadership feel like status while still making it meaningful?",
        "Which parts go into the Strategic Plan, and which parts stay internal to the Leadership OS?",
    ]:
        bullet(doc, item)

    h1(doc, "Final Recommendation")
    para(
        doc,
        "Keep the ambition. Replace the corporate costume. The club should operate like a serious technical team: clear ownership, visible artifacts, short feedback loops, documented support, and role transition when a function is not being performed. That is strong. Calling it a company with KPIs and PIPs is weaker than the underlying idea and easier for people to attack.",
    )


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = COLORS["ink"]
    for style_name in ["List Bullet", "List Number"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(10)

    add_title_page(doc)
    add_super_prompt(doc)
    add_audit(doc)
    add_corrected_framework(doc)
    add_role_scorecards(doc)
    add_improvement_template(doc)
    add_timeline_and_language(doc)
    add_ben_review(doc)

    core = doc.core_properties
    core.title = "AI Club Leadership Accountability Audit and Rewrite"
    core.subject = "Audit and corrected leadership operating system for Northern Highlands AI Club"
    core.author = "Codex"
    core.keywords = "AI Club, leadership, accountability, June 1, operating system"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
